"""Generate the project report as an editable Word .docx (python-docx).

Produces Ramco-Agent-Report.docx with real Word headings, tables, and bullet
lists (so it reflows cleanly and is easy to edit further). Metrics are the ones
measured by scripts/evaluate.py.

Run with:  python -m scripts.make_report_docx
"""
from __future__ import annotations

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.shared import Pt, Inches, RGBColor

ACCENT = RGBColor(0x0B, 0x5C, 0xAB)
OUT = "Ramco-Agent-Report.docx"


def _set_base_style(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10.5)
    for s in doc.sections:
        s.top_margin = s.bottom_margin = Inches(0.6)
        s.left_margin = s.right_margin = Inches(0.7)


def _meta_table(doc: Document, pairs: list[tuple[str, str]]) -> None:
    t = doc.add_table(rows=0, cols=2)
    t.allow_autofit = True
    for label, value in pairs:
        cells = t.add_row().cells
        r = cells[0].paragraphs[0].add_run(label)
        r.bold = True
        r.font.size = Pt(9.5)
        rv = cells[1].paragraphs[0].add_run(value)
        rv.font.size = Pt(9.5)
    for row in t.rows:
        for c in row.cells:
            c.paragraphs[0].paragraph_format.space_after = Pt(1)


def _data_table(doc: Document, headers: list[str], rows: list[list[str]],
                right_cols: set[int] | None = None) -> None:
    right_cols = right_cols or set()
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Light Grid Accent 1"
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    for i, h in enumerate(headers):
        run = t.rows[0].cells[i].paragraphs[0].add_run(h)
        run.bold = True
        run.font.size = Pt(9.5)
    for r in rows:
        cells = t.add_row().cells
        for i, val in enumerate(r):
            p = cells[i].paragraphs[0]
            run = p.add_run(val)
            run.font.size = Pt(9.5)
            if i in right_cols:
                from docx.enum.text import WD_ALIGN_PARAGRAPH
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT


def _diagram(doc: Document, lines: list[str]) -> None:
    tbl = doc.add_table(rows=1, cols=1)
    tbl.style = "Table Grid"
    cell = tbl.rows[0].cells[0]
    cell.paragraphs[0].text = ""
    for ln in lines:
        p = cell.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.space_before = Pt(0)
        run = p.add_run(ln)
        run.font.name = "Consolas"
        run.font.size = Pt(7.5)
    # drop the empty leading paragraph
    cell.paragraphs[0]._element.getparent().remove(cell.paragraphs[0]._element)


def build() -> str:
    doc = Document()
    _set_base_style(doc)

    title = doc.add_heading("Real-World Context-Engineered Auto-Reply Agent", level=0)
    sub = doc.add_paragraph()
    sr = sub.add_run("A confidence-gated LLM support agent for Ramco products, delivered as a Telegram bot")
    sr.italic = True
    sr.font.size = Pt(11)
    sr.font.color.rgb = RGBColor(0x5B, 0x64, 0x70)

    _meta_table(doc, [
        ("Author", "Nikhil Kumar Jha"),
        ("Date", "4 June 2026"),
        ("Platform", "Telegram bot (@ramco_assist_bot)"),
        ("Knowledge domain", "Ramco products (public information)"),
        ("Hosted dashboard", "https://caring-gentleness-production-c3e7.up.railway.app"),
        ("Deployment", "Railway — poller + dashboard + Postgres"),
        ("Models", "Claude Haiku 4.5 (router) · Claude Sonnet 4.6 (drafter)"),
        ("Repository", "github.com/mrxjha/lumenx-agent — /Real-world-reply-agent"),
    ])

    doc.add_heading("1 · Overview", level=1)
    doc.add_paragraph(
        "The agent answers questions about Ramco's enterprise software (ERP, HCM & Global Payroll, "
        "Aviation/MRO, Logistics, EAM) on a real messaging platform. An incoming Telegram message is "
        "classified by intent, answered from a curated knowledge base with Claude, scored by a small "
        "neural network, and then auto-sent when confidence is high or routed to a human-review dashboard "
        "when it is not. Every model call is cost-logged. The objective is to automate routine replies "
        "while never hallucinating on pricing or contract terms, keeping a human in the loop until the "
        "model earns trust."
    )

    doc.add_heading("2 · System Architecture", level=1)
    _diagram(doc, [
        "Telegram msg -> Intent Router (Haiku) -> greeting/off-topic -> polite reply",
        "                    | pricing - refund - technical - other",
        "                    v",
        "               Context Builder",
        "                 - LLM Wiki (Ramco product pages + company policy)",
        "                 - current chat history + past-conversation gist",
        "                 - feedback log (past approved Q->A pairs)",
        "                    v",
        "               Drafter (Sonnet) - anti-hallucination, cites sources, logs cost",
        "                    v",
        "               Confidence Net (Tiny MLP) -> score 0-1",
        "                    | >=0.90 & AUTO_SEND on        < 0.90 (or AUTO_SEND off)",
        "                    v                              v",
        "                auto-send                   human-review dashboard -> approve -> send",
        "                    \\--------------> feedback log -> labels for next MLP retrain",
    ])
    p = doc.add_paragraph()
    pr = p.add_run("The pipeline is platform-agnostic — the Telegram connector is the only messaging-specific "
                   "module, so an email/Teams backend could be swapped in without touching the agent logic.")
    pr.italic = True
    pr.font.size = Pt(9)

    doc.add_heading("3 · Build Process", level=1)
    for b in [
        "LLM Wiki: five Ramco product lines + a company-policy page rendered to structured Markdown with a "
        "cross-reference graph (Karpathy “LLM-wiki” pattern). Pricing is recorded as “not published "
        "publicly,” so correct deflection becomes the grounded answer.",
        "Cold-start data & Confidence Net: a bot cannot scrape past chats, so 47 realistic "
        "(question -> draft -> ideal-reply) examples bootstrap the MLP — negatives reproduce feared failures "
        "(invented prices, missing citations); positives are grounded and cited.",
        "Pipeline: intent router (cheap model) -> context builder -> drafter (quality model) with an "
        "anti-hallucination prompt -> per-call token/cost logging.",
        "Dashboard: Streamlit review queue with confidence badge, expandable context window, cost panel, and "
        "one-click approve/edit/reject that replies via Telegram.",
        "Deployment: one Docker image, three roles (poller / dashboard / web) on Railway with managed Postgres.",
    ]:
        doc.add_paragraph(b, style="List Bullet")

    doc.add_heading("4 · Context Engineering", level=1)
    doc.add_heading("4.1 Intent router", level=2)
    doc.add_paragraph(
        "A cheap Haiku call labels each message (greeting · pricing · refund · technical · other) and returns "
        "JSON with its own confidence. Greetings and off-topic chat are answered politely without invoking the "
        "expensive drafter, and the label biases retrieval toward the right wiki pages.")
    doc.add_heading("4.2 The context window", level=2)
    doc.add_paragraph(
        "For each reply the builder assembles four ranked sources into one prompt: (1) the LLM Wiki — company "
        "policy plus the top pages by keyword overlap with the query, expanded along the cross-reference graph; "
        "(2) the current chat history; (3) a summary of past conversations; and (4) a feedback log of previously "
        "approved Q->A pairs for the same intent. The dashboard can show this exact window per reply for auditability.")
    doc.add_heading("4.3 Anti-hallucination guardrail", level=2)
    doc.add_paragraph(
        "The drafter's system prompt forbids inventing prices, refund windows, or contract terms; if the answer "
        "is not in the loaded wiki it must say so and offer to route to sales. Facts are cited inline with source "
        "tags (e.g. [ramco_hcm], [company_policy]), which also feed the confidence model.")
    doc.add_heading("4.4 Confidence Net (the routing brain)", level=2)
    doc.add_paragraph(
        "A small scikit-learn MLP (hidden layers 32->16) scores each draft 0-1 — “would a human send this "
        "as-is?” — using only signals available before sending: draft length/word-count, digit count "
        "(pricing-risk), citation count, hedge-phrase presence, sign-off compliance, retrieval-hit count, and the "
        "intent one-hot. Labels come from edit-distance between the draft and the human-sent reply "
        "(near-duplicate -> 1; heavily edited/rejected -> 0) plus explicit thumbs feedback. Drafts scoring at or "
        "above a configurable threshold (default 0.90) qualify for auto-send; a global AUTO_SEND_ENABLED switch "
        "keeps everything in human review during cold-start.")

    doc.add_heading("5 · Deployment", level=1)
    doc.add_paragraph(
        "One Docker image is dispatched into three roles via a SERVICE_ROLE variable on Railway, all sharing a "
        "managed Postgres database:")
    _data_table(doc,
        ["Service", "Role", "Public?", "Function"],
        [
            ["poller", "poller", "No", "Telegram getUpdates loop -> pipeline -> route"],
            ["dashboard", "dashboard", "Yes — the hosted URL", "Human review queue + cost dashboard"],
            ["Postgres", "—", "No", "Shared state: threads, drafts, feedback, token usage"],
        ])
    note = doc.add_paragraph()
    nr = note.add_run(
        "Engineering notes: the corporate network blocks Telegram, so the live bot is driven from a mobile "
        "device while the build/test loop runs locally against the (reachable) model API — Railway runs outside "
        "that firewall. Deploy fixes: a Streamlit health-check path (so traffic isn't routed before the app "
        "binds), a Postgres connect-timeout, and suppression of HTTP-client logs that would otherwise print the "
        "bot token.")
    nr.italic = True
    nr.font.size = Pt(9)

    doc.add_heading("6 · Performance Analysis", level=1)
    doc.add_paragraph(
        "Measured on a balanced 10-query set (pricing, refund, technical, other, greeting) run through the full "
        "production pipeline against the live models.")
    _data_table(doc,
        ["Metric", "Result"],
        [
            ["Intent accuracy", "100% (10 / 10)"],
            ["Avg reply latency", "4.30 s (median 4.34 s; range 3.47–5.25 s)"],
            ["Cost per reply", "$0.0099 (~$9.86 per 1,000 replies)"],
            ["True hallucination rate", "0%"],
        ])

    doc.add_heading("6.1 Accuracy & Speed", level=2)
    doc.add_paragraph(
        "The intent router labelled 10 / 10 queries correctly; on grounded drafts the Confidence Net scored "
        "0.99–1.00. Drafting (Sonnet) dominates latency, while routing (Haiku) is sub-second; a “typing…” "
        "indicator is shown while drafting.")

    doc.add_heading("6.2 Cost", level=2)
    _data_table(doc,
        ["Component", "Model", "Avg / reply", "Share"],
        [
            ["Intent routing", "Claude Haiku 4.5", "$0.00058", "6%"],
            ["Reply drafting", "Claude Sonnet 4.6", "$0.00928", "94%"],
            ["Confidence scoring", "local MLP", "$0.00000", "0%"],
            ["Total", "", "$0.00986", "100%"],
        ], right_cols={2, 3})
    cp = doc.add_paragraph()
    cpr = cp.add_run("The cheap-router / quality-drafter split keeps classification effectively free; local MLP "
                     "scoring adds zero marginal cost.")
    cpr.italic = True
    cpr.font.size = Pt(9)

    doc.add_heading("6.3 Hallucination rate", level=2)
    doc.add_paragraph(
        "Across the 4 sensitive (pricing/refund) queries the agent quoted 0 fabricated prices or terms — all four "
        "deflected to a custom quote / the signed contract, with citations. An automated keyword check flagged 1 "
        "reply, but manual review showed a false positive: it described the pricing model (“per employee, per "
        "country”) while stating there is no published fee — no number was invented. True hallucination rate: 0%.")

    doc.add_heading("6.4 Confidence model", level=2)
    doc.add_paragraph(
        "Bootstrapped on 47 generated examples (29 positive / 18 negative). It fits this small distribution "
        "cleanly; a held-out split and weekly retraining on real dashboard feedback are the path to a calibrated "
        "threshold.")

    doc.add_heading("7 · Limitations & Next Steps", level=1)
    for b in [
        "The MLP is bootstrapped on synthetic data; real dashboard approve/edit/reject feedback will replace it "
        "and enable a proper held-out evaluation and threshold calibration.",
        "Wiki relevance uses keyword overlap; embeddings would improve retrieval on paraphrased questions.",
        "Hallucination is measured by heuristic + manual review on a small set; an LLM-graded rubric on a larger "
        "held-out set would tighten the estimate.",
        "Prompt caching on the reused wiki/context block would cut drafting cost further.",
    ]:
        doc.add_paragraph(b, style="List Bullet")

    doc.add_heading("Appendix — Screenshots", level=1)
    doc.add_paragraph(
        "[ Insert here: (1) the dashboard review queue showing a confidence score and the expandable context "
        "window, and (2) the Telegram chat showing a pricing deflection. ]")

    doc.save(OUT)
    return OUT


if __name__ == "__main__":
    path = build()
    print(f"wrote {path}")
